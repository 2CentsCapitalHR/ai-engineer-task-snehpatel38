import os
import json
from io import BytesIO
from typing import List, Dict, Tuple
import streamlit as st
from docx import Document
from PyPDF2 import PdfReader
import numpy as np
from tqdm import tqdm

# LLM client (Gemini)
try:
    import google.generativeai as genai
except Exception:
    genai = None

# Try to import faiss; if unavailable, fall back to sklearn cosine similarity
USE_FAISS = True
try:
    import faiss
except Exception:
    USE_FAISS = False
    from sklearn.metrics.pairwise import cosine_similarity

# ---------------------------
# ========== CONFIG =========
# ---------------------------
EMBED_MODEL_DEFAULT = "models/embedding-001"  # free-tier embedding model
LLM_MODEL_DEFAULT = "gemini-1.5-flash"        # free-tier text model

CHECKLISTS = {
    "Company Incorporation": [
        "Articles of Association",
        "Memorandum of Association",
        "Board Resolution",
        "Register of Members and Directors",
        "UBO Declaration Form"
    ],
    "Employment Contract": [
        "Standard Employment Contract",
        "Job Description",
        "Signed Offer Letter"
    ],
    "Data Protection Policy": [
        "Appropriate Policy Document",
        "Privacy Notice"
    ]
}

DOC_TYPE_KEYWORDS = {
    "Articles of Association": ["articles of association", "aoa", "articles"],
    "Memorandum of Association": ["memorandum of association", "moa", "memorandum"],
    "Board Resolution": ["board resolution", "resolution for incorporation", "resolution"],
    "Register of Members and Directors": ["register of members", "register of directors"],
    "UBO Declaration Form": ["ubo", "ultimate beneficial owner", "ubo declaration"],
    "Standard Employment Contract": ["employment contract", "standard employment contract"],
    "Appropriate Policy Document": ["appropriate policy document", "data protection", "privacy policy"]
}

# ---------------------------
# ========== HELPERS =========
# ---------------------------
def init_gemini(api_key_env='GEMINI_API_KEY'):
    key = os.environ.get(api_key_env) or (st.secrets.get(api_key_env) if hasattr(st, "secrets") else None)
    if key is None:
        st.warning(f"Set your Gemini API key in environment variable {api_key_env} or Streamlit secrets.")
        return None
    if genai is None:
        st.error("google-generativeai library is not installed.")
        return None
    genai.configure(api_key=key)
    return genai

def extract_docx_paragraphs(file_bytes: bytes) -> Tuple[List[Dict], Document]:
    doc = Document(BytesIO(file_bytes))
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            paragraphs.append({"text": text, "paragraph_index": i})
    return paragraphs, doc

def extract_pdf_text(path_or_bytes) -> str:
    text = ""
    if isinstance(path_or_bytes, (bytes, bytearray)):
        reader = PdfReader(BytesIO(path_or_bytes))
    else:
        reader = PdfReader(path_or_bytes)
    for p in reader.pages:
        page_text = p.extract_text()
        if page_text:
            text += "\n" + page_text
    return text

def chunk_text(text: str, chunk_size=800, overlap=100) -> List[str]:
    tokens = text.split()
    chunks = []
    i = 0
    while i < len(tokens):
        chunk = tokens[i:i+chunk_size]
        chunks.append(" ".join(chunk))
        i += chunk_size - overlap
    return chunks

def embed_texts(genai_client, texts: List[str], model=EMBED_MODEL_DEFAULT) -> np.ndarray:
    if genai_client is None:
        raise ValueError("Gemini client not initialized.")
    embs = []
    for txt in texts:
        resp = genai.embed_content(model=model, content=txt)
        embs.append(resp["embedding"])
    return np.array(embs).astype("float32")

def build_index(embeddings: np.ndarray):
    if USE_FAISS:
        d = embeddings.shape[1]
        index = faiss.IndexFlatIP(d)
        norms = np.linalg.norm(embeddings, axis=1, keepdims=True)
        norms[norms == 0] = 1.0
        emb_norm = embeddings / norms
        index.add(emb_norm)
        return {"type": "faiss", "index": index, "emb_norm": emb_norm}
    else:
        return {"type": "sklearn", "embeddings": embeddings}

def retrieve_top_k(index_struct, query_emb: np.ndarray, k=4):
    if index_struct["type"] == "faiss":
        qn = query_emb / (np.linalg.norm(query_emb) + 1e-10)
        D, I = index_struct["index"].search(np.array([qn]).astype("float32"), k)
        return I[0].tolist(), D[0].tolist()
    else:
        sims = cosine_similarity(query_emb.reshape(1, -1), index_struct["embeddings"])[0]
        idxs = np.argsort(-sims)[:k]
        return idxs.tolist(), sims[idxs].tolist()

def detect_documents(uploaded_files: List[Tuple[str, bytes]]) -> Dict[str, List[str]]:
    detected = {}
    for fname, content in uploaded_files:
        lower = fname.lower()
        assigned = None
        for dtype, keywords in DOC_TYPE_KEYWORDS.items():
            if any(kw in lower for kw in keywords):
                assigned = dtype
                break
        if not assigned:
            try:
                paras, _ = extract_docx_paragraphs(content)
                sample_text = " ".join([p["text"] for p in paras[:2]]).lower()
                for dtype, keywords in DOC_TYPE_KEYWORDS.items():
                    if any(kw in sample_text for kw in keywords):
                        assigned = dtype
                        break
            except Exception:
                assigned = "Unknown Document"
        if assigned is None:
            assigned = "Unknown Document"
        detected.setdefault(assigned, []).append(fname)
    return detected

def check_paragraph_with_gemini(genai_client, paragraph_text: str, retrieved_contexts: str, model=LLM_MODEL_DEFAULT):
    prompt = f"""
You are an ADGM compliance assistant. Given the user's document paragraph and the retrieved ADGM reference snippets, identify any compliance issues.
Return a JSON array with:
- section
- issue
- severity
- suggestion

Paragraph:
\"\"\"{paragraph_text}\"\"\"

ADGM Reference:
\"\"\"{retrieved_contexts}\"\"\"

Respond ONLY with valid JSON array.
"""
    try:
        model_obj = genai.GenerativeModel(model)
        resp = model_obj.generate_content(prompt)
        out_text = resp.text
    except Exception as e:
        st.error(f"Gemini generation error: {e}")
        return []
    try:
        return json.loads(out_text)
    except:
        import re
        m = re.search(r"\[.*\]", out_text, flags=re.S)
        return json.loads(m.group(0)) if m else []

def insert_inline_notes(doc: Document, findings_by_paragraph: Dict[int, List[Dict]]):
    for p_idx, issues in findings_by_paragraph.items():
        if not issues:
            continue
        if p_idx < len(doc.paragraphs):
            para = doc.paragraphs[p_idx]
            notes = [f"{iss.get('severity','?')}: {iss.get('issue','')}. Suggest: {iss.get('suggestion','')}" for iss in issues]
            para.add_run(" [REVIEW_NOTE: " + " | ".join(notes) + " ]")

# ---------------------------
# ========== STREAMLIT UI =========
# ---------------------------
st.set_page_config(page_title="ADGM Corporate Agent", layout="wide")
st.title("ADGM Corporate Agent — Demo (Gemini + RAG + Checklist)")

st.sidebar.header("Settings")
emb_model = st.sidebar.text_input("Embedding model", value=EMBED_MODEL_DEFAULT)
llm_model = st.sidebar.text_input("LLM model", value=LLM_MODEL_DEFAULT)
top_k = st.sidebar.number_input("Top-k retrieval", min_value=1, max_value=8, value=3)
chunk_size = st.sidebar.number_input("Ref chunk size (words)", min_value=200, max_value=2000, value=800, step=100)
overlap = st.sidebar.number_input("Ref chunk overlap", min_value=20, max_value=500, value=100, step=10)

genai_client = init_gemini()

st.header("1) ADGM Reference")
uploaded_ref = st.file_uploader("Upload ADGM reference PDF", type=["pdf"])
adgm_text = ""
if uploaded_ref:
    adgm_text = extract_pdf_text(uploaded_ref.read())
else:
    if os.path.exists("Data Sources.pdf"):
        adgm_text = extract_pdf_text("Data Sources.pdf")

if adgm_text:
    if "adgm_chunks" not in st.session_state:
        chunks = chunk_text(adgm_text, chunk_size=chunk_size, overlap=overlap)
        st.session_state["adgm_chunks"] = chunks
    else:
        chunks = st.session_state["adgm_chunks"]

    if "adgm_index" not in st.session_state and genai_client:
        with st.spinner("Creating embeddings..."):
            emb = embed_texts(genai_client, chunks, model=emb_model)
            idx_struct = build_index(emb)
            st.session_state["adgm_index"] = idx_struct
            st.session_state["adgm_embs"] = emb

st.header("2) Upload .docx files")
uploaded_docs = st.file_uploader("Upload one or more .docx files", type=["docx"], accept_multiple_files=True)

if uploaded_docs:
    files_for_detection = [(f.name, f.read()) for f in uploaded_docs]
    detected_map = detect_documents(files_for_detection)
    st.subheader("Detected document types")
    st.json(detected_map)

    inferred_process = None
    all_detected_types = list(detected_map.keys())
    if any(k in ["Articles of Association", "Memorandum of Association"] for k in all_detected_types):
        inferred_process = "Company Incorporation"
    elif any("Employment" in k for k in all_detected_types):
        inferred_process = "Employment Contract"

    if inferred_process:
        required = CHECKLISTS.get(inferred_process, [])
        present = []
        recognized = set(detected_map.keys())
        for req in required:
            if req in recognized:
                present.append(req)
        missing = [r for r in required if r not in present]
        checklist_report = {
            "process": inferred_process,
            "required_documents": len(required),
            "present_documents": present,
            "missing_documents": missing
        }
        st.subheader("Checklist report")
        st.json(checklist_report)

    overall_results = []
    for fname, content in files_for_detection:
        paragraphs, doc = extract_docx_paragraphs(content)
        findings_by_paragraph = {}
        for p in tqdm(paragraphs, desc=f"Checking {fname}"):
            text = p["text"]
            pidx = p["paragraph_index"]
            retrieved_text = ""
            if "adgm_index" in st.session_state and genai_client:
                q_emb = embed_texts(genai_client, [text], model=emb_model)[0]
                I, _ = retrieve_top_k(st.session_state["adgm_index"], q_emb, k=top_k)
                retrieved_chunks = [st.session_state["adgm_chunks"][i] for i in I]
                retrieved_text = "\n---\n".join(retrieved_chunks)
            issues = []
            if genai_client:
                issues = check_paragraph_with_gemini(genai_client, text, retrieved_text, model=llm_model)
            findings_by_paragraph[pidx] = issues
        insert_inline_notes(doc, findings_by_paragraph)
        out_buf = BytesIO()
        doc.save(out_buf)
        out_buf.seek(0)
        report = {
            "filename": fname,
            "inferred_process": inferred_process,
            "findings": [
                {"paragraph_index": pi, "issues": issues} for pi, issues in findings_by_paragraph.items() if issues
            ]
        }
        overall_results.append({
            "filename": fname,
            "reviewed_docx_bytes": out_buf.getvalue(),
            "report": report
        })

    st.markdown("### Download outputs")
    for item in overall_results:
        st.download_button(f"Download reviewed {item['filename']}", data=item["reviewed_docx_bytes"],
                           file_name=f"reviewed_{item['filename']}",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        st.download_button(f"Download JSON report {item['filename']}", data=json.dumps(item["report"], indent=2),
                           file_name=f"report_{item['filename']}.json", mime="application/json")